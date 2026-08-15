#!/usr/bin/env python3
"""Render a syllabus quick reference into the formats a class actually needs.

Usage:
    python render.py QUICKREF.md [--html OUT.html] [--docx OUT.docx] [--columns 2]

Reads a deliberately small slice of Markdown: h1, h2, bullets, pipe tables,
bold spans, and > for the closing note. A quick reference should not need
anything more expressive than that, and a rewrite that wants a feature this does
not support is too complicated for a one-pager.

No third-party Markdown dependency on purpose: this runs on a school laptop
with whatever Python happens to be installed.
"""
import argparse
import html
import re

# bookSHelf house theme: parchment canvas, one Wedgwood accent held under 5% of
# the surface, warm ink ramp, serif headings capped at weight 500, no italic.
# Full spec: bookSHelf/.claude/skills/theme-factory/themes/bookshelf.md
CSS = """
@page { size: letter; margin: 0.45in; }
:root {
  --wedgwood: #4e6e8e; --wedgwood-deep: #3d5a80;
  --parchment: #f5f4ed; --warm-sand: #e8e6dc;
  --ink: #141413; --charcoal: #3d3d3a; --olive: #504e49; --stone: #6b6a64;
  --border-cream: #f0eee6; --border-warm: #e8e6dc;
  --serif: "Source Serif 4", "Source Serif Pro", Charter, Georgia, "Times New Roman", serif;
  --sans: Inter, system-ui, -apple-system, "Segoe UI", Helvetica, Arial, sans-serif;
}
* { box-sizing: border-box; }
body { font: 10.5pt/1.55 var(--sans); color: var(--ink);
       background: var(--parchment); margin: 0; padding: 0.45in; max-width: 8.5in;
       -webkit-print-color-adjust: exact; print-color-adjust: exact; }
h1 { font: 500 25pt/1.1 var(--serif); margin: 0 0 3pt; letter-spacing: -0.4px; }
.sub { font-size: 10pt; color: var(--olive); margin: 0 0 10pt;
       padding-bottom: 7pt; border-bottom: 1px solid var(--wedgwood); }
.cols { column-count: __COLS__; column-gap: 26pt; counter-reset: sec; }
/* Sections never split across the column boundary: a rule you are scanning for
   should be whole and under its own heading, not continued overleaf. That makes
   section length the thing that has to be managed. Keep each one to roughly
   four or five bullets and the columns pack evenly on their own. One nine-item
   section cannot pack, and it is the section that wants dividing anyway. */
section { break-inside: avoid; margin: 0 0 10pt; counter-increment: sec; }
h2 { font: 500 13pt/1.2 var(--serif); color: var(--ink); margin: 0 0 5pt;
     padding-bottom: 3pt; border-bottom: 1px solid var(--border-warm);
     break-after: avoid; break-inside: avoid; }
h2::before { content: counter(sec, decimal-leading-zero);
             display: block; font: 500 9pt/1.4 var(--serif);
             color: var(--wedgwood); letter-spacing: 0.5px; }
p { margin: 0 0 5pt; color: var(--charcoal); }
ul { margin: 0; padding: 0; list-style: none; }
li { position: relative; padding-left: 13pt; margin-bottom: 2.5pt;
     color: var(--charcoal); break-inside: avoid; }
table, tr { break-inside: avoid; }
li::before { content: "\\2013"; position: absolute; left: 0;
             color: var(--wedgwood); }
table { border-collapse: collapse; width: 100%; margin: 2pt 0 4pt;
        font-variant-numeric: tabular-nums; }
th { text-align: left; font: 500 8pt/1.4 var(--sans); text-transform: uppercase;
     letter-spacing: 1.2px; color: var(--charcoal); background: var(--warm-sand);
     padding: 2pt 6pt; }
td { padding: 2pt 6pt; border-bottom: 1px solid var(--border-cream);
     color: var(--charcoal); }
strong { font-weight: 600; color: var(--ink); }
a { color: var(--wedgwood-deep); text-decoration: none; }
.note { column-span: all; margin-top: 6pt; padding-top: 5pt;
        border-top: 1px solid var(--border-warm); font-size: 9pt;
        color: var(--stone); }
@media print { body { padding: 0; } .note { page-break-inside: avoid; } }
"""


def inline(s):
    s = html.escape(s)
    s = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", s)
    s = re.sub(r"(?<![\">])(https?://[^\s<)]+)", r'<a href="\1">\1</a>', s)
    return s


def parse(md):
    """-> (title, subtitle, [(heading, [blocks])], [notes]).

    A block is ('ul', [items]) | ('table', [rows]) | ('p', text).
    """
    title, sub, notes, sections = "", "", [], []
    cur = None
    lines = md.splitlines()
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        i += 1
        if not line:
            continue
        if line.startswith("# "):
            title = line[2:].strip()
            for j in range(i, len(lines)):          # first non-blank line after the
                nxt = lines[j].strip()              # title is the who/where/when strip
                if not nxt:
                    continue
                if not nxt.startswith(("#", "-", "*", "|", ">")):
                    sub, i = nxt, j + 1
                break
            continue
        if line.startswith("## "):
            cur = (line[3:].strip(), [])
            sections.append(cur)
            continue
        if line.startswith(">"):
            notes.append(line.lstrip("> ").strip())
            continue
        if cur is None:
            cur = ("", [])
            sections.append(cur)
        if line.startswith(("- ", "* ")):
            if cur[1] and cur[1][-1][0] == "ul":
                cur[1][-1][1].append(line[2:].strip())
            else:
                cur[1].append(("ul", [line[2:].strip()]))
        elif line.startswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            if all(re.fullmatch(r":?-{2,}:?", c) for c in cells if c):
                continue                                  # separator row
            if cur[1] and cur[1][-1][0] == "table":
                cur[1][-1][1].append(cells)
            else:
                cur[1].append(("table", [cells]))
        else:
            cur[1].append(("p", line))
    return title, sub, sections, notes


def to_html(md, columns):
    title, sub, sections, notes = parse(md)
    out = [CSS.replace("__COLS__", str(columns))]
    body = ['<h1>%s</h1>' % inline(title)]
    if sub:
        body.append('<p class="sub">%s</p>' % inline(sub))
    body.append('<div class="cols">')
    for head, blocks in sections:
        body.append("<section>")
        if head:
            body.append("<h2>%s</h2>" % inline(head))
        for kind, payload in blocks:
            if kind == "ul":
                body.append("<ul>%s</ul>" %
                            "".join("<li>%s</li>" % inline(x) for x in payload))
            elif kind == "table":
                head_row, *rest = payload
                body.append("<table><tr>%s</tr>%s</table>" % (
                    "".join("<th>%s</th>" % inline(c) for c in head_row),
                    "".join("<tr>%s</tr>" % "".join("<td>%s</td>" % inline(c) for c in r)
                            for r in rest)))
            else:
                body.append("<p>%s</p>" % inline(payload))
        body.append("</section>")
    body.append("</div>")
    for n in notes:
        body.append('<p class="note">%s</p>' % inline(n))
    return ("<!doctype html><html><head><meta charset=\"utf-8\">"
            "<title>%s</title><style>%s</style></head><body>%s</body></html>"
            % (html.escape(title), out[0], "".join(body)))


def to_docx(md, path):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    WEDGWOOD = RGBColor(0x4E, 0x6E, 0x8E)
    INK = RGBColor(0x14, 0x14, 0x13)
    CHARCOAL = RGBColor(0x3D, 0x3D, 0x3A)
    OLIVE = RGBColor(0x50, 0x4E, 0x49)
    STONE = RGBColor(0x6B, 0x6A, 0x64)
    SERIF = "Georgia"       # local stand-in for Source Serif 4, same warm figure
    SANS = "Segoe UI"       # local stand-in for Inter

    title, sub, sections, notes = parse(md)
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.5)
        s.left_margin = s.right_margin = Inches(0.6)
    normal = doc.styles["Normal"]
    normal.font.name = SANS
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = CHARCOAL
    normal.paragraph_format.line_spacing = 1.55
    normal.paragraph_format.space_after = Pt(3)

    def shade(cell, hex_fill):
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        el = OxmlElement("w:shd")
        el.set(qn("w:val"), "clear")
        el.set(qn("w:fill"), hex_fill)
        cell._tc.get_or_add_tcPr().append(el)

    def plain(text):
        """Render **bold** spans into a fresh paragraph and return it."""
        p = doc.add_paragraph()
        for k, part in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
            if part:
                p.add_run(part).bold = bool(k % 2)
        return p

    h = doc.add_paragraph()
    r = h.add_run(title)
    r.font.name, r.font.size, r.font.color.rgb = SERIF, Pt(22), INK
    h.paragraph_format.line_spacing = 1.1
    h.paragraph_format.space_after = Pt(2)
    if sub:
        p = doc.add_paragraph()
        r = p.add_run(sub)
        r.font.size, r.font.color.rgb = Pt(10), OLIVE
        p.paragraph_format.space_after = Pt(10)

    for n, (head, blocks) in enumerate(sections, 1):
        if head:
            p = doc.add_paragraph()          # 01 / 02 section marker, Wedgwood
            r = p.add_run("%02d" % n)
            r.font.name, r.font.size, r.font.color.rgb = SERIF, Pt(8.5), WEDGWOOD
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(8), Pt(0)
            p = doc.add_paragraph()
            r = p.add_run(head)
            r.font.name, r.font.size, r.font.color.rgb = SERIF, Pt(12), INK
            p.paragraph_format.line_spacing = 1.2
            p.paragraph_format.space_after = Pt(2)
        for kind, payload in blocks:
            if kind == "ul":
                for item in payload:
                    # En-dash bullets in Wedgwood, per the theme. Word's own
                    # bullet glyph cannot be recolored independently of the text,
                    # so the marker is a literal run instead of a list style.
                    p = doc.add_paragraph()
                    r = p.add_run("–  ")
                    r.font.color.rgb = WEDGWOOD
                    for k, part in enumerate(re.split(r"\*\*(.+?)\*\*", item)):
                        if part:
                            run = p.add_run(part)
                            run.bold = bool(k % 2)
                            if run.bold:
                                run.font.color.rgb = INK
                    p.paragraph_format.left_indent = Inches(0.16)
                    p.paragraph_format.first_line_indent = Inches(-0.16)
                    p.paragraph_format.space_after = Pt(1)
            elif kind == "table":
                head_row, *rest = payload
                t = doc.add_table(rows=1, cols=len(head_row))
                t.style = "Table Grid"
                for c, txt in zip(t.rows[0].cells, head_row):
                    shade(c, "E8E6DC")                     # warm sand header
                    c.text = ""
                    run = c.paragraphs[0].add_run(txt.upper())
                    run.font.size, run.font.color.rgb = Pt(8), CHARCOAL
                for row in rest:
                    cells = t.add_row().cells
                    for c, txt in zip(cells, row):
                        c.text = ""
                        run = c.paragraphs[0].add_run(re.sub(r"\*\*", "", txt))
                        run.font.size, run.font.color.rgb = Pt(10), CHARCOAL
            else:
                plain(payload)

    for n in notes:
        p = plain(n)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(8)
        for run in p.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = STONE
    doc.save(path)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("quickref")
    ap.add_argument("--html")
    ap.add_argument("--docx")
    ap.add_argument("--columns", type=int, default=2)
    a = ap.parse_args()
    md = open(a.quickref, encoding="utf-8").read()
    if not a.html and not a.docx:
        a.html = a.quickref.rsplit(".", 1)[0] + ".html"
    if a.html:
        open(a.html, "w", encoding="utf-8").write(to_html(md, a.columns))
        print("[OK] wrote " + a.html)
    if a.docx:
        to_docx(md, a.docx)
        print("[OK] wrote " + a.docx)


if __name__ == "__main__":
    main()
