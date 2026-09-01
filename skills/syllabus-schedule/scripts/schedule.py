#!/usr/bin/env python3
"""Pull the course schedule out of a syllabus .docx and print it.

    python schedule.py SYLLABUS.docx [--theme NAME] [--out DIR] [--name STEM]

Writes STEM-schedule.md, .html and .pdf. The Markdown is the editable artifact:
re-run against the same .md (pass it instead of the .docx) after hand edits.

A schedule table is found by its header row, not its position, because the shape
differs per course: Integrated Math 1 numbers lessons (#, Ch, Assignment, Date),
Statistics splits odd and even rotation days into two date columns, and Computer
Science breaks the year into fifteen per-chapter tables. Calendar tables (Date,
Weekday, ...) come along too. Everything else in the syllabus is left behind.

Rendering borrows the syllabus-quickref renderer and its themes, so a schedule
and a quick reference for the same course look like siblings.
"""
import argparse, os, subprocess, sys

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

QUICKREF = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(
    os.path.abspath(__file__)))), "syllabus-quickref", "scripts")

# Two rules the quickref renderer cannot carry, injected after it runs.
#
# The compartment box is set break-inside:avoid, which is right for a one-page
# reference and wrong here: a semester of lessons never fits one page, so the
# rule only strands the first page half empty before Chrome gives up and splits
# anyway. Schedules fragment.
#
# ponytail: :has() is Chrome-only, and Chrome is the only engine that prints
# these pages. The renderer's Markdown slice cannot mark a row, so the shading
# goes in afterwards: a no-school row is the one with no lesson number.
SHADE = ("<style>section{break-inside:auto}"
         "tbody tr:has(td:first-child:empty){background:var(--block);"
         "color:var(--olive);font-style:italic}</style>")


def is_schedule(hdr):
    """A schedule or calendar table, told apart from the policy tables."""
    return bool(hdr) and (hdr[0] in ("#", "Unit")
                          or (hdr[0] == "Date" and "Weekday" in hdr))


def as_md(table):
    rows = [[c.text.strip().replace("|", r"\|") for c in r.cells] for r in table.rows]
    sep = "|" + "|".join(["---"] * len(rows[0])) + "|"
    return "\n".join(["| " + " | ".join(rows[0]) + " |", sep]
                     + ["| " + " | ".join(r) + " |" for r in rows[1:]])


def harvest(path):
    """-> (title, byline, note, [(heading, intro, table_md), ...])."""
    doc = Document(path)
    lead = [p.text.strip() for p in doc.paragraphs if p.text.strip()][:2]
    info = {}
    for t in doc.tables:
        for r in t.rows:
            cells = [c.text.strip() for c in r.cells]
            if len(cells) == 2:
                info.setdefault(cells[0], cells[1])

    sections, heads, para, note, last_para = [], {}, "", "", ""
    for child in doc.element.body.iterchildren():
        if child.tag.endswith("}p"):
            p = Paragraph(child, doc)
            if not p.text.strip():
                continue
            level = p.style.name
            # Heading 1 and 2 name the document and the section it lives in, so
            # they would prefix every compartment with the course's own title.
            if level.startswith("Heading") and level[-1] not in "12":
                heads[level] = p.text.strip()
                # A new heading at this level retires anything deeper.
                for deeper in [k for k in heads if k > level]:
                    del heads[deeper]
                para = ""
            elif not level.startswith("Heading"):
                para = last_para = p.text.strip()
        elif child.tag.endswith("}tbl"):
            table = Table(child, doc)
            hdr = [c.text.strip() for c in table.rows[0].cells]
            if is_schedule(hdr):
                title = " · ".join(heads[k] for k in sorted(heads))
                # The paragraph before the first table becomes the closing note,
                # unless it is already this section's own intro line.
                if not sections and last_para != para:
                    note = last_para
                sections.append((title or "Schedule", para, as_md(table)))
                heads.clear()
            para = ""

    if not sections:
        sys.exit("[FAIL] no schedule tables in %s. Headers looked for: "
                 "'#', 'Unit', or 'Date' + 'Weekday'." % os.path.basename(path))

    title = (lead[0] if lead else "Course") + ": Course Schedule"
    byline = " · ".join(v for v in (info.get("Instructor"),
                                    info.get("Courses and Time"),
                                    lead[1] if len(lead) > 1 else None) if v)
    return title, byline, note, sections


def to_markdown(path):
    title, byline, note, sections = harvest(path)
    out = ["# " + title, ""]
    if byline:
        out += [byline, ""]
    for head, para, table in sections:
        out += ["## " + head, ""]
        if para:
            out += ["- " + para, ""]
        out += [table, ""]
    if note:
        out += ["> " + note, ""]
    return "\n".join(out) + "\n"


def render(md_path, theme, html_path, pdf_path):
    subprocess.run([sys.executable, os.path.join(QUICKREF, "render.py"), md_path,
                    "--theme", theme, "--columns", "1", "--html", html_path],
                   check=True)
    html = open(html_path, encoding="utf-8").read()
    open(html_path, "w", encoding="utf-8").write(html.replace("</head>", SHADE + "</head>", 1))
    sys.path.insert(0, QUICKREF)
    from render import to_pdf
    to_pdf(html_path, pdf_path)


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("syllabus", help="the syllabus .docx, or an already-built -schedule.md")
    ap.add_argument("--theme", default="annotated", help="any syllabus-quickref theme")
    ap.add_argument("--out", default=".", help="directory for the .html and .pdf")
    ap.add_argument("--name", help="output stem (default: taken from the input filename)")
    a = ap.parse_args()

    os.makedirs(a.out, exist_ok=True)
    if a.syllabus.lower().endswith(".md"):
        md_path = a.syllabus
        stem = a.name or os.path.splitext(os.path.basename(md_path))[0]
    else:
        stem = a.name or os.path.splitext(os.path.basename(a.syllabus))[0] + "-schedule"
        md_path = os.path.join(a.out, stem + ".md")
        open(md_path, "w", encoding="utf-8").write(to_markdown(a.syllabus))
        print("[OK] wrote %s" % md_path, flush=True)

    html = os.path.join(a.out, stem + ".html")
    render(md_path, a.theme, html, os.path.join(a.out, stem + ".pdf"))
    print("[OK] wrote %s and %s" % (html, os.path.join(a.out, stem + ".pdf")))


if __name__ == "__main__":
    main()
